from .unitutil.file import absjoin, test_file_dir
from docx.api import Document


dir_pkg_path = absjoin(test_file_dir, 'expanded_docx')
python_powered_path = absjoin(test_file_dir, 'python-powered.png')


class DescribeFloatingShapeAdd(object):
    def it_has_part_as_header_part(self):
        document = Document(dir_pkg_path)
        header = document.add_header()
        paragraph = header.add_paragraph()
        run = paragraph.add_run()
        run.add_floating_picture(python_powered_path)
        document.save('/home/daniel/faboozle.docx')
